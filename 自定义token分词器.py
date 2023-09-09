import nltk
from nltk.tokenize import word_tokenize
from read_clean不带接口 import read_and_clean_pdf_text
from docx import Document
import 按token切割段落
from 按token切割段落 import breakdown_txt_to_satisfy_token_limit_for_pdf

def split_text(text, max_tokens):
    # 对文本进行分词
    tokens = word_tokenize(text)

    # 初始化段落列表和当前段落
    paragraphs = []
    current_paragraph = []

    for token in tokens:
        # 如果当前段落的token数目已经达到最大限制，则将其添加到段落列表中，并开始新的段落
        if len(current_paragraph) >= max_tokens:
            paragraphs.append(current_paragraph)
            current_paragraph = []

        # 将当前token添加到当前段落中
        current_paragraph.append(token)

    # 将最后一个段落添加到段落列表中
    paragraphs.append(current_paragraph)

    # 将每个段落中的tokens连接成字符串
    paragraphs = [' '.join(paragraph) for paragraph in paragraphs]

    return paragraphs


def list_to_word(input_list, output_file):
    doc = Document()
    for item in input_list:
        if item is not None:
            doc.add_paragraph(str(item))
        else:
            doc.add_paragraph("None")

    doc.save(output_file)

path = r'C:\Users\Administrator\Desktop\PDF\ChatGPT赋能图书馆知识服务：原理、场景与进路_郭亚军.pdf'
# txt = breakdown_txt_to_satisfy_token_limit_for_pdf(read_and_clean_pdf_text(path)[0], get_token_fn=count_tokens,limit=1500)
txt = split_text(read_and_clean_pdf_text(path)[0],600)
output_file_path = r"C:\Users\Administrator\Desktop\output_cut.docx"
list_to_word(txt,output_file_path)